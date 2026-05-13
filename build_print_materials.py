from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(45, 71, 102)
MUTED = RGBColor(85, 97, 113)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_margins(doc):
    for section in doc.sections:
      section.top_margin = Inches(0.7)
      section.bottom_margin = Inches(0.7)
      section.left_margin = Inches(0.7)
      section.right_margin = Inches(0.7)


def base_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].font.color.rgb = RGBColor(18, 25, 36)

    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = ACCENT

    for level in [1, 2, 3]:
        name = f"Heading {level}"
        styles[name].font.name = "Aptos Display"
        styles[name].font.bold = True
        styles[name].font.color.rgb = ACCENT
        styles[name].paragraph_format.space_before = Pt(10)
        styles[name].paragraph_format.space_after = Pt(5)


def add_title_block(doc, eyebrow, title, subtitle):
    p = doc.add_paragraph()
    r = p.add_run(eyebrow.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph(title, style="Title")
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(subtitle)
    p.paragraph_format.space_after = Pt(12)


def add_cover_page(doc, eyebrow, title, subtitle, image_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(eyebrow.upper())
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = "Aptos Display"
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(18, 25, 36)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(image_path, width=Inches(6.3))
    except Exception:
        fallback = p.add_run("[Immagine di copertina non disponibile]")
        fallback.italic = True
        fallback.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("4 squadre  |  3 round  |  60 minuti  |  1 anno simulato")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT

    doc.add_page_break()


def add_info_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        set_cell_shading(row[0], "EEF3F8")
    doc.add_paragraph()
    return table


def add_round_sheet(doc, round_title):
    doc.add_heading(round_title, level=2)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    prompts = [
        ("Decisione strategica scelta", ""),
        ("Decisione operativa scelta", ""),
        ("Coinvolgimento interno scelto", ""),
        ("Evento straordinario emerso", ""),
        ("Perché il team ha scelto questa linea", ""),
        ("Quale prezzo siete disposti a pagare", ""),
    ]
    for left, right in prompts:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        set_cell_shading(row[0], "EEF3F8")
    doc.add_paragraph()


def build_participant_doc(path):
    doc = Document()
    set_margins(doc)
    base_styles(doc)
    add_cover_page(
        doc,
        "ItalBridge Transport",
        "Materiale partecipanti",
        "Business game executive | Trasporti e logistica",
        "benvenuto.png",
    )
    add_title_block(
        doc,
        "ItalBridge Transport",
        "Materiale partecipanti",
        "Scheda di lavoro stampabile per il business game executive.",
    )

    add_info_table(
        doc,
        [
            ("Durata", "60 minuti"),
            ("Formato", "4 squadre | 3 round | 1 anno simulato"),
            ("Cosa osservare", "Clienti, marginalità, cassa, servizio, tempo strategico"),
            ("Mandato della squadra", "........................................................"),
        ],
    )

    doc.add_heading("Il business game", level=1)
    for text in [
        "La simulazione riproduce un anno di decisioni in un operatore integrato di trasporti e logistica.",
        "Ogni squadra parte dallo stesso scenario iniziale, ma affronta il percorso con un mandato riservato assegnato dal facilitatore.",
        "In ogni round il team prende una decisione strategica, una decisione operativa e una decisione sul coinvolgimento interno del management.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Scenario di partenza", level=1)
    for text in [
        "ItalBridge Transport parte con ricavi solidi, ma con marginalità contenuta, cassa sotto pressione e struttura operativa già molto tesa.",
        "Il mercato resta competitivo, con clienti rilevanti, richieste di flessibilità e costi esterni instabili.",
        "Durante il game possono entrare eventi straordinari interni ed esterni che modificano priorità, tempi e qualità della risposta manageriale.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Regole del gioco", level=1)
    for text in [
        "Discutete nel team e comunicate al facilitatore una scelta per ciascun blocco del round.",
        "Non esiste una scelta perfetta in assoluto: conta la coerenza tra priorità protetta, costo pagato e robustezza della traiettoria.",
        "Usate questa scheda solo come supporto di lavoro: annotate ciò che vi aiuta davvero a decidere e a leggere il debrief finale.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Uso rapido", level=1)
    for text in [
        "Annotate solo ciò che vi aiuta a chiarire la scelta.",
        "Per ogni round fissate prima la priorità che volete proteggere.",
        "Segnate il prezzo che siete disposti a pagare per quella linea.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    for round_name in ["Round 1", "Round 2", "Round 3"]:
        add_round_sheet(doc, round_name)

    doc.add_heading("Debrief finale", level=1)
    for text in [
        "Che cosa abbiamo privilegiato davvero nel nostro percorso?",
        "Quale prezzo abbiamo pagato per proteggere quel risultato?",
        "Da quale punto forte o fragile ripartiremmo se l’anno continuasse?",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.save(path)


def build_facilitator_doc(path):
    doc = Document()
    set_margins(doc)
    base_styles(doc)
    add_title_block(
        doc,
        "ItalBridge Transport",
        "Guida facilitatore",
        "Uso operativo del business game in aula.",
    )

    add_info_table(
        doc,
        [
            ("Pagine da usare", "index.html | settings.html | game.html | epilogo.html"),
            ("Durata consigliata", "60 minuti"),
            ("Formato", "13 direttori | 4 squadre | facilitatore unico"),
            ("Obiettivo", "Far emergere trade-off, costi pagati e qualità della traiettoria"),
        ],
    )

    doc.add_heading("Checklist prima dell’aula", level=1)
    for text in [
        "Aprire settings.html e assegnare i mandati alle 4 squadre.",
        "Salvare i mandati e verificare che risultino completi.",
        "Aprire index.html sullo schermo condiviso come pagina di benvenuto.",
        "Tenere pronto game.html per avviare il game quando il gruppo è allineato.",
        "Verificare che demo finale e debrief finale si aprano correttamente.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Checklist operativa del giorno del workshop", level=1)
    for text in [
        "Aprire la pagina di benvenuto sullo schermo condiviso prima dell’ingresso dei partecipanti.",
        "Impostare i mandati in settings.html sullo stesso browser che verrà usato per il game.",
        "Tenere aperta la guida facilitatore per consultare scaletta e domande di debrief.",
        "Avviare il game solo dopo aver verificato che i mandati risultino salvati.",
        "Nel debrief finale, usare prima la lettura trasversale e solo dopo il confronto squadra per squadra.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Scaletta da 60 minuti", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Fase"
    hdr[1].text = "Tempo"
    hdr[2].text = "Indicazione per il facilitatore"
    for c in hdr:
        set_cell_shading(c, "DCE6F2")
    rows = [
        ("Apertura e contesto", "5-7 min", "Mostrare index.html e spiegare scenario, logica dei 3 round e regola del debrief."),
        ("Round 1", "12-15 min", "Far discutere le squadre e poi inserire una squadra alla volta."),
        ("Round 2", "12-15 min", "Tenere alto il ritmo: prima decisione, poi evento, poi chiusura round."),
        ("Round 3", "12-15 min", "Chiudere tutte le squadre e aprire la lettura finale."),
        ("Debrief", "10-15 min", "Usare la pagina finale per confrontare risultato, costo pagato e traiettoria futura."),
    ]
    for fase, tempo, nota in rows:
        row = table.add_row().cells
        row[0].text = fase
        row[1].text = tempo
        row[2].text = nota
    doc.add_paragraph()

    doc.add_heading("Messaggi chiave da passare", level=1)
    for text in [
        "Non cercate la scelta perfetta, cercate una scelta coerente con il vostro mandato e con il contesto.",
        "Il tempo strategico conta quanto i numeri: non tutto si risolve intervenendo di più.",
        "Il debrief non serve a trovare vincitori e vinti, ma a leggere qualità della traiettoria e costo pagato.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Domande guida per il debrief", level=1)
    for text in [
        "Chi ha centrato il proprio mandato con il prezzo più alto?",
        "Chi ha protetto meglio il sistema nel suo insieme?",
        "Chi appare più forte oggi ma più fragile domani?",
        "Quale squadra ha costruito la traiettoria più robusta per l’anno successivo?",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Se qualcosa si inceppa", level=1)
    for text in [
        "Se una squadra rallenta, chiedere di esplicitare in una frase la priorità che sta cercando di proteggere.",
        "Se il gruppo si perde nei dettagli, riportare il confronto su margine, cassa, servizio, tempo strategico.",
        "Se serve mostrare il finale senza rigiocare tutto, usare i pulsanti demo già presenti nel game.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.save(path)


if __name__ == "__main__":
    build_participant_doc("Materiale partecipanti - ItalBridge Transport.docx")
    build_facilitator_doc("Guida facilitatore - ItalBridge Transport.docx")
